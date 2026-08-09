from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "0B1220"
TEAL = "14B8A6"
SLATE = "475569"
PALE = "E6FFFB"
LIGHT = "F1F5F9"
RED = "B42318"
AMBER = "B54708"
GREEN = "067647"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("RAKSHEX  •  CONTROLLED RELEASE  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "RAKSHEX  /  PRODUCTION LAUNCH & OPERATIONS"
    header.style = "Caption"
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.style = "Title"
    p.add_run(title)
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.style = "Subtitle"


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.add_run(f"{index}.  ").bold = True
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)
        if len(table.rows) % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text, color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE if color == TEAL else LIGHT)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    add_header_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.2)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Subtitle"].font.name = "Aptos"
    styles["Subtitle"].font.size = Pt(13)
    styles["Subtitle"].font.color.rgb = RGBColor.from_string(SLATE)
    for name, size in (("Heading 1", 19), ("Heading 2", 12)):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor.from_string(NAVY if name == "Heading 1" else TEAL)
        styles[name].paragraph_format.space_before = Pt(8)
        styles[name].paragraph_format.space_after = Pt(5)


def build(output: Path):
    doc = Document()
    configure(doc)

    # 1 — Cover and control
    add_title(
        doc,
        "Production Launch\n& Operations Binder",
        "Rakshex release 2026-07-30  |  Deployable code + controlled launch evidence",
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "RELEASE POSITION",
        "Engineering gates pass locally. Public paid launch remains controlled by live-provider, "
        "operational, financial, and legal approvals listed in this binder.",
    )
    add_table(
        doc,
        ["Document control", "Value"],
        [
            ("Product", "Rakshex AI agent and API security platform"),
            ("Release", "2026-07-30 / source package + VSIX 0.2.1"),
            ("Classification", "Internal launch control; legal drafts are unexecuted"),
            ("Prepared for", "Business owner, Engineering, Security, Operations, Finance, Counsel"),
            ("Change rule", "Record deployed commit, evidence links, owner, date, and approver"),
        ],
        [1.7, 5.5],
    )
    add_body(
        doc,
        "Do not replace blank company, officer, tax, jurisdiction, customer, price, or signature "
        "fields with assumptions. Complete them from authoritative records and obtain qualified review.",
    )

    # 2 — Decision dashboard
    page_break(doc)
    add_h1(doc, "1. Release decision dashboard")
    add_table(
        doc,
        ["Gate", "Local state", "Promotion requirement"],
        [
            ("Typecheck / build", "PASS", "Repeat on exact release commit in protected CI"),
            ("Automated tests", "PASS", "874 passed; run live DB + Playwright gates in CI"),
            ("Dependency security", "PASS", "Lock policy, peers, and production audit clean"),
            ("Web + API contract", "PASS", "Verify production origins, CORS, and health endpoints"),
            ("VS Code extension", "PASS", "Install VSIX; Marketplace publisher verification"),
            ("Worker / queues", "READY", "Deploy worker separately and prove queue drain"),
            ("Email", "CONTROLLED", "Verified domain + SMTP delivery exercises"),
            ("Billing", "CONTROLLED", "Real payment, failure, refund, webhook replay, reconciliation"),
            ("Legal / tax", "OPEN", "Complete drafts, GST/tax process, counsel approval"),
        ],
        [1.55, 1.0, 4.65],
    )
    add_callout(
        doc,
        "GO / NO-GO RULE",
        "Private beta may proceed only after worker, SMTP, production origins, monitoring, and the "
        "signed HTTPS buyer journey pass. Paid GA additionally requires billing, tax, legal, and on-call sign-off.",
        AMBER,
    )
    add_h2(doc, "Unsupported launch claims")
    add_bullets(
        doc,
        [
            "Do not claim certifications, patents, audits, market-first status, uptime, or benchmarks without evidence.",
            "Do not publish customer names, logos, quotations, or usage statistics without written authorization.",
            "The compliance engine maps evidence; it does not award SOC 2, ISO, GDPR, or regulatory certification.",
        ],
    )

    # 3 — Product and architecture
    page_break(doc)
    add_h1(doc, "2. Product scope and service topology")
    add_body(
        doc,
        "The simple user journey is: create account → select workspace → import an API collection → "
        "run a scan → understand findings → assign/remediate → export evidence. Advanced tools remain "
        "under “More tools” so first-time users are not forced through enterprise controls.",
    )
    add_table(
        doc,
        ["Component", "Responsibility", "Production dependency"],
        [
            ("Web", "Marketing, auth, dashboard, team, billing, findings", "Vercel + API origin"),
            ("API", "Auth, tenancy, scans, findings, billing, integrations", "Railway + Postgres + Redis"),
            ("Worker", "Scan and notification queue processing", "Separate Railway service + Redis"),
            ("PostgreSQL", "Authoritative tenant and billing state", "Migrations + backup/restore"),
            ("Redis", "Queues, cache, distributed rate limits", "Required; no production fallback"),
            ("VS Code", "Editor findings, scan actions, control-plane links", "API key + production API"),
        ],
        [1.2, 3.3, 2.7],
    )
    add_h2(doc, "Tenant and billing guarantees")
    add_bullets(
        doc,
        [
            "Workspace permission checks guard collections, findings, team, keys, and billing operations.",
            "Workspace subscriptions enforce seats; roles include owner, admin, developer/editor, analyst/auditor, and viewer.",
            "Money is stored as integer minor units; provider event IDs support idempotent webhook processing.",
            "API keys are prefixed for environment clarity, displayed once, and stored as hashes.",
        ],
    )

    # 4 — Deployment
    page_break(doc)
    add_h1(doc, "3. Production deployment order")
    add_numbered(
        doc,
        [
            "Freeze the release commit; require green CI, secret scan, dependency audit, container scan, and SBOM.",
            "Create a recoverable PostgreSQL backup and run migrations from the release exactly once.",
            "Deploy API with production JWT, encryption, database, Redis, CORS, SMTP, OAuth, monitoring, and provider secrets.",
            "Deploy the worker from railway.worker.toml; enqueue a canary scan and prove it completes.",
            "Deploy web with the exact API, site, WebSocket, and Sentry public configuration.",
            "Install the VSIX in a clean VS Code profile; sign in, health-check, import, scan, refresh, and open a finding.",
            "Run the HTTPS buyer journey and cross-tenant denial probe; attach timestamps and screenshots.",
            "Enable billing only after payment/refund/reconciliation and legal/tax approvals are signed.",
        ],
    )
    add_h2(doc, "Production configuration groups")
    add_table(
        doc,
        ["Group", "Examples", "Rule"],
        [
            ("Core", "NODE_ENV, DATABASE_URL, REDIS_URL, JWT_SECRET, ENCRYPTION_KEY", "Secret manager only"),
            ("Origins", "APP_URL, FRONTEND_URL, CORS_ORIGINS, NEXT_PUBLIC_*", "Exact HTTPS origins"),
            ("Email", "SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASS, SMTP_FROM", "Verified domain"),
            ("Payments", "Stripe/Razorpay keys, webhook secrets, price IDs", "Server-only; live test"),
            ("Identity", "Google/GitHub OAuth, GitHub App", "Production callback URLs"),
            ("Observability", "Sentry DSN, OTLP, uptime, paging", "Named owner + alert test"),
        ],
        [1.1, 3.55, 2.55],
    )

    # 5 — Team subscriptions
    page_break(doc)
    add_h1(doc, "4. Team subscriptions and customer operations")
    add_table(
        doc,
        ["Journey", "Expected behavior", "Evidence"],
        [
            ("Create workspace", "Creator becomes owner; workspace selector persists", "Workspace ID + owner"),
            ("Choose plan/seats", "Quote uses integer minor units; minimum seats enforced", "Provider order + DB row"),
            ("Invite teammate", "Duplicate member/invite blocked; email delivered", "Invite + delivery log"),
            ("Accept invite", "Correct workspace/role; no cross-tenant access", "Member row + 403 probe"),
            ("Change role", "Authorized roles only; last owner protected", "Audit event"),
            ("Seat limit", "New member blocked at limit with clear upgrade path", "UI/API response"),
            ("Cancel", "Provider and local status agree; access follows policy", "Webhook + reconciliation"),
        ],
        [1.35, 3.7, 2.15],
    )
    add_h2(doc, "Support language")
    add_bullets(
        doc,
        [
            "Tell the user what happened, what is safe, and the single next action.",
            "Never expose provider errors, stack traces, tenant IDs, secrets, or raw payment payloads.",
            "For pending webhooks, say the plan is being confirmed and provide a refresh/support path.",
            "Do not promise refunds, SLA credits, certifications, or legal outcomes outside approved policy.",
        ],
    )

    # 6 — UX
    page_break(doc)
    add_h1(doc, "5. User-friendly product standard")
    add_table(
        doc,
        ["Surface", "First-time user must understand", "Required check"],
        [
            ("Navigation", "Overview, import, findings, scans, team, billing", "Primary tasks visible; advanced collapsed"),
            ("Empty states", "Why the page is empty and one next action", "No dead ends or mystery IDs"),
            ("Errors", "What failed, whether data is safe, and retry/support", "No endpoint or HTML parse leakage"),
            ("Progress", "Queued, running, completed, failed", "Refresh/polling behaves consistently"),
            ("Permissions", "Why an action is unavailable", "Role-aware buttons and server enforcement"),
            ("Billing", "Plan, seats, price unit, renewal/cancel state", "Provider and local state match"),
            ("VS Code", "Sign in → import/scan → findings → dashboard", "Commands and status bar use same terms"),
        ],
        [1.2, 3.55, 2.45],
    )
    add_h2(doc, "Accessibility and responsive release check")
    add_bullets(
        doc,
        [
            "Keyboard-only navigation, visible focus, logical heading order, labeled form controls, and error summaries.",
            "Contrast and state are not communicated by color alone; loading and destructive actions are explicit.",
            "Verify 390 px mobile, tablet, desktop, zoom at 200%, reduced motion, and screen-reader announcements.",
            "Test new-account, viewer, developer, admin, owner, invited-user, expired-invite, and over-seat states.",
        ],
    )

    # 7 — Security/ops
    page_break(doc)
    add_h1(doc, "6. Security, reliability, and incident readiness")
    add_table(
        doc,
        ["Control", "Release implementation", "Production exercise"],
        [
            ("Fail closed", "Redis/SMTP/provider requirements enforced by environment", "Remove dependency; verify safe failure"),
            ("Secrets", "Hashing/encryption and secret scanning", "Rotate test credential; inspect logs"),
            ("Tenancy", "Workspace authorization and denial tests", "Cross-tenant 403 probe"),
            ("Queues", "API enqueue + independent worker", "Canary, retry, dead-letter/alert"),
            ("Recovery", "Backup/restore scripts and runbooks", "Timed restore to clean environment"),
            ("Monitoring", "Health/readiness, Sentry, telemetry hooks", "Trigger and acknowledge alert"),
            ("Supply chain", "Frozen lock, age policy, audit, SBOM", "Verify exact release artifact"),
        ],
        [1.25, 3.45, 2.5],
    )
    add_h2(doc, "Incident first hour")
    add_numbered(
        doc,
        [
            "Declare severity and incident commander; preserve timestamps and evidence.",
            "Contain: revoke keys, pause billing/webhooks, disable integration, or activate kill switch as appropriate.",
            "Assess tenants and data; do not speculate in customer communications.",
            "Notify required internal, contractual, insurer, and regulatory contacts using approved timelines.",
            "Recover from known-good state; monitor; record decisions; schedule root-cause review.",
        ],
    )

    # 8 — Legal
    page_break(doc)
    add_h1(doc, "7. Legal and business paperwork control")
    add_callout(
        doc,
        "UNEXECUTED DRAFTS",
        "The five supplied Word files are included unchanged. They must not be published or signed until "
        "the fields below are completed consistently and qualified counsel approves the intended jurisdictions.",
        RED,
    )
    add_table(
        doc,
        ["Draft", "Open fields observed", "Required owners"],
        [
            ("Grievance Officer notice", "12", "Business / HR / counsel"),
            ("Cookie Policy", "7", "Privacy / web / counsel"),
            ("Privacy Policy", "9", "Privacy / security / counsel"),
            ("Terms of Service", "14", "Business / finance / counsel"),
            ("Enterprise MSA + Order", "32", "Sales / finance / security / counsel"),
        ],
        [2.6, 1.45, 3.15],
    )
    add_bullets(
        doc,
        [
            "Entity name, registration number, GSTIN, registered address, trading-name relationship, and signatory.",
            "Named Grievance Officer, tested privacy/security contacts, effective dates, governing law, and notices.",
            "Cookie inventory and consent behavior; subprocessors, retention, security, and data-request operations.",
            "Fees, taxes, currency, renewal/cancellation, SLA, liability, insurance, DPA, and customer order details.",
        ],
    )

    # 9 — Signoff
    page_break(doc)
    add_h1(doc, "8. Launch sign-off record")
    add_table(
        doc,
        ["Gate", "Owner", "Evidence link / reference", "Decision + date"],
        [
            ("Exact-commit CI", "Engineering", "", ""),
            ("Backup / restore", "Engineering", "", ""),
            ("API + worker canary", "Engineering", "", ""),
            ("SMTP delivery", "Operations", "", ""),
            ("Buyer journey", "Product", "", ""),
            ("Payments / refund", "Finance", "", ""),
            ("Security tabletop", "Security", "", ""),
            ("Privacy exercise", "Privacy owner", "", ""),
            ("Legal / tax", "Counsel + Finance", "", ""),
            ("Claims review", "Business owner", "", ""),
        ],
        [1.55, 1.35, 2.85, 1.45],
    )
    add_body(doc, "Engineering approver: __________________  Date: __________  Commit: __________________")
    add_body(doc, "Security / Privacy: _____________________  Date: __________")
    add_body(doc, "Finance / Operations: ___________________  Date: __________")
    add_body(doc, "Business owner: _________________________  Date: __________")
    add_body(doc, "Qualified counsel: ______________________  Date: __________")

    # 10 — Launch day
    page_break(doc)
    add_h1(doc, "9. Launch-day command sheet")
    add_h2(doc, "Before traffic")
    add_bullets(
        doc,
        [
            "Confirm release SHA, migrations, API ready check, worker queue drain, web origin, status page, and alert routing.",
            "Confirm support, security, privacy, grievance, and billing inboxes are staffed and tested.",
            "Confirm public copy matches the approved claims register and the current legal documents.",
        ],
    )
    add_h2(doc, "First buyer journey")
    add_numbered(
        doc,
        [
            "Register and verify email; create and switch workspace.",
            "Invite a second user; accept; verify role and tenant isolation.",
            "Import Postman/OpenAPI; scan; observe progress; open, assign, resolve, and export a finding.",
            "Install VSIX; authenticate; health-check; refresh findings; open dashboard.",
            "Buy team seats; add member; verify invoice/payment state; test cancellation/refund only in approved environment.",
        ],
    )
    add_h2(doc, "Rollback triggers")
    add_bullets(
        doc,
        [
            "Cross-tenant access, secret exposure, incorrect charges, migrations affecting data integrity, or queue loss.",
            "Widespread auth/invite failure, inability to revoke access, or monitoring blind spot during incident.",
            "Public legal/claims mismatch that could materially mislead a buyer.",
        ],
    )
    add_callout(
        doc,
        "FINAL RULE",
        "A green build makes the artifact deployable. Recorded operational evidence and accountable approvals make the launch authorized.",
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: build_launch_binder.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
